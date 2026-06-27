#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Génère le cahier des charges Assigame — version ~14 pages."""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

DOCS_DIR = Path(__file__).resolve().parent
ESGIS_LOGO = DOCS_DIR / "assets" / "esgis-logo.png"
COVER_PAGE_HEIGHT = Cm(27.2)
COVER_PAGE_WIDTH = Cm(18.6)
sys.path.insert(0, str(DOCS_DIR / "content"))

from metadata import COVER, METADATA, MOTS_CLES, RESUME  # noqa: E402
from sections_data import (  # noqa: E402
    ACTEURS,
    API_ENDPOINTS_COMPACT,
    CAS_UTILISATION_COMPACT,
    DIAGRAMS,
    EXIGENCES_FONCTIONNELLES_COMPACT,
    EXIGENCES_NON_FONCTIONNELLES,
    SCENARIOS_TEST,
    SCREENSHOTS,
    STACK_TECHNIQUE,
)

OUTPUT = DOCS_DIR / "Rapport du projet assigame.docx"
DIAGRAMS_DIR = DOCS_DIR / "diagrams"
SCREENSHOTS_DIR = DOCS_DIR / "screenshots"
TARGET_PAGES = 14


def setup_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.12
    normal.paragraph_format.space_after = Pt(4)

    for level, size in [(1, 14), (2, 12), (3, 11)]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)


def centered(doc: Document, text: str, size: int = 11, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold


def text(doc: Document, content: str) -> None:
    doc.add_paragraph(content)


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for p in t.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()


def figure(doc: Document, path: Path, caption: str, width: float = 5.0) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(caption)
    r.bold = True
    r.font.size = Pt(10)
    if path.is_file():
        img = doc.add_paragraph()
        img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img.add_run().add_picture(str(path), width=Inches(width))


def figures_side_by_side(doc: Document, left: tuple[Path, str], right: tuple[Path, str], width: float = 3.1) -> None:
    t = doc.add_table(rows=1, cols=2)
    for i, (path, cap) in enumerate([left, right]):
        cell = t.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(cap + "\n").bold = True
        if path.is_file():
            p.add_run().add_picture(str(path), width=Inches(width))
    doc.add_paragraph()


def screenshots_pair(doc: Document, pair: list[tuple[str, str, str]], width: float = 3.2) -> None:
    t = doc.add_table(rows=1, cols=2)
    for i, (fid, fname, title) in enumerate(pair):
        cell = t.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"Figure {fid} — {title}\n").bold = True
        path = SCREENSHOTS_DIR / fname
        if path.is_file():
            p.add_run().add_picture(str(path), width=Inches(width))
    doc.add_paragraph()


def set_table_frame(table, border_sz: str = "18") -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)

    for existing in tbl_pr.findall(qn("w:tblW")):
        tbl_pr.remove(existing)

    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(COVER_PAGE_WIDTH.twips))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.insert(0, tbl_w)

    for existing in tbl_pr.findall(qn("w:tblLayout")):
        tbl_pr.remove(existing)
    tbl_layout = OxmlElement("w:tblLayout")
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_pr.append(tbl_layout)

    for existing in tbl_pr.findall(qn("w:jc")):
        tbl_pr.remove(existing)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    tbl_pr.append(jc)

    grid = tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for grid_col in grid.findall(qn("w:gridCol")):
            grid_col.set(qn("w:w"), str(COVER_PAGE_WIDTH.twips))

    for existing in tbl_pr.findall(qn("w:tblBorders")):
        tbl_pr.remove(existing)
    tbl_borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), border_sz)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")
        tbl_borders.append(element)
    for edge in ("insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        tbl_borders.append(element)
    tbl_pr.append(tbl_borders)


def set_cell_padding(cell, margin_dxa: int = 260) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for side in ("top", "left", "bottom", "right"):
        margin = OxmlElement(f"w:{side}")
        margin.set(qn("w:w"), str(margin_dxa))
        margin.set(qn("w:type"), "dxa")
        tc_mar.append(margin)
    tc_pr.append(tc_mar)


def clear_page_border(section) -> None:
    pg_borders = section._sectPr.find(qn("w:pgBorders"))
    if pg_borders is not None:
        section._sectPr.remove(pg_borders)


def set_page_border(section, border_sz: str = "18", space: str = "6") -> None:
    clear_page_border(section)
    pg_borders = OxmlElement("w:pgBorders")
    pg_borders.set(qn("w:offsetFrom"), "text")
    pg_borders.set("{http://schemas.microsoft.com/office/word/2010/wordml}display", "allPages")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), border_sz)
        border.set(qn("w:space"), space)
        border.set(qn("w:color"), "000000")
        pg_borders.append(border)
    section._sectPr.append(pg_borders)


def remove_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    tbl_borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        tbl_borders.append(element)
    tbl_pr.append(tbl_borders)


def add_paragraph_bottom_line(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def cover_run(
    paragraph,
    text: str,
    *,
    size: int = 11,
    bold: bool = False,
    italic: bool = False,
) -> None:
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)


def cover_cell_lines(cell, lines: list[str], align: WD_ALIGN_PARAGRAPH, *, underline: bool = False) -> None:
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after = Pt(1)
        cover_run(p, line, size=10, bold=(i == 0))
    if underline:
        line_p = cell.add_paragraph()
        line_p.alignment = align
        line_p.paragraph_format.space_before = Pt(2)
        line_p.paragraph_format.space_after = Pt(0)
        add_paragraph_bottom_line(line_p)


def cover_spacer(cell, height_pt: int) -> None:
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(height_pt)
    cover_run(p, "", size=2)


def cover_line_cell(
    cell,
    text: str,
    *,
    size: int = 11,
    bold: bool = False,
    italic: bool = False,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER,
    space_before: int = 0,
    space_after: int = 6,
    first: bool = False,
) -> None:
    p = cell.paragraphs[0] if first else cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    cover_run(p, text, size=size, bold=bold, italic=italic)


def cover(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Cm(1.0)
    sec.bottom_margin = Cm(1.0)
    sec.left_margin = Cm(1.15)
    sec.right_margin = Cm(1.15)
    clear_page_border(sec)

    frame = doc.add_table(rows=1, cols=1)
    set_table_frame(frame, border_sz="18")
    frame.rows[0].height = COVER_PAGE_HEIGHT
    frame.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    frame_cell = frame.rows[0].cells[0]
    frame_cell.width = COVER_PAGE_WIDTH
    set_cell_padding(frame_cell)

    intro = frame_cell.paragraphs[0]
    intro.paragraph_format.space_before = Pt(0)
    intro.paragraph_format.space_after = Pt(0)
    intro.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    intro.paragraph_format.line_spacing = Pt(1)

    header = frame_cell.add_table(rows=1, cols=2)
    remove_table_borders(header)
    header.rows[0].cells[0].width = Cm(8.2)
    header.rows[0].cells[1].width = Cm(8.2)
    cover_cell_lines(header.rows[0].cells[0], COVER["ministry_lines"], WD_ALIGN_PARAGRAPH.LEFT, underline=True)
    cover_cell_lines(header.rows[0].cells[1], COVER["republic_lines"], WD_ALIGN_PARAGRAPH.RIGHT, underline=True)

    cover_spacer(frame_cell, 10)
    cover_line_cell(frame_cell, COVER["mser"], size=10, bold=True, space_after=0)
    cover_spacer(frame_cell, 28)

    if ESGIS_LOGO.is_file():
        logo = frame_cell.add_paragraph()
        logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo.paragraph_format.space_before = Pt(0)
        logo.paragraph_format.space_after = Pt(4)
        logo.add_run().add_picture(str(ESGIS_LOGO), width=Inches(5.1))

    cover_line_cell(frame_cell, COVER["school"], size=10, space_after=0)
    cover_spacer(frame_cell, 42)

    for line in COVER["document_type"]:
        cover_line_cell(frame_cell, line, size=15, bold=True, space_after=2)

    cover_line_cell(frame_cell, COVER["project_title"], size=12, italic=True, space_before=8, space_after=0)
    cover_spacer(frame_cell, 70)

    footer = frame_cell.add_table(rows=1, cols=2)
    remove_table_borders(footer)
    footer.rows[0].cells[0].width = Cm(8.2)
    footer.rows[0].cells[1].width = Cm(8.2)
    left = footer.rows[0].cells[0]
    right = footer.rows[0].cells[1]

    group_label = left.paragraphs[0]
    group_label.alignment = WD_ALIGN_PARAGRAPH.LEFT
    group_label.paragraph_format.space_after = Pt(4)
    cover_run(group_label, COVER["group_label"], size=11)

    for member in COVER["group_members"]:
        member_p = left.add_paragraph()
        member_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        member_p.paragraph_format.left_indent = Cm(0.6)
        member_p.paragraph_format.space_after = Pt(2)
        cover_run(member_p, f"- {member}", size=11)

    supervisor_p = right.paragraphs[0]
    supervisor_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    supervisor_p.paragraph_format.space_before = Pt(18)
    supervisor_p.paragraph_format.space_after = Pt(4)
    cover_run(supervisor_p, f"{COVER['supervisor_prefix']} ", size=11)
    cover_run(supervisor_p, COVER["supervisor"], size=11, bold=True)

    teacher_p = right.add_paragraph()
    teacher_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    cover_run(teacher_p, COVER["teacher_note"], size=11)

    cover_spacer(frame_cell, 36)
    cover_line_cell(
        frame_cell,
        f"{COVER['academic_year_label']} {COVER['academic_year']}",
        size=11,
        space_after=0,
    )

    doc.add_page_break()

    body_sec = doc.add_section()
    clear_page_border(body_sec)
    body_sec.top_margin = Cm(2.2)
    body_sec.bottom_margin = Cm(2.2)
    body_sec.left_margin = Cm(2.2)
    body_sec.right_margin = Cm(2.2)


def fix_heading_colors(docx_path: Path) -> None:
    """Force les grands titres en noir dans styles.xml et stylesWithEffects.xml."""
    import shutil
    import tempfile
    import xml.etree.ElementTree as ET
    import zipfile

    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    heading_ids = {
        "Heading1",
        "Heading2",
        "Heading3",
        "Heading1Char",
        "Heading2Char",
        "Heading3Char",
        "Title",
        "TOCHeading",
    }

    def patch(xml_bytes: bytes) -> bytes:
        root = ET.fromstring(xml_bytes)
        for style in root.findall("w:style", ns):
            sid = style.get(f"{W}styleId", "")
            if sid not in heading_ids:
                continue
            r_pr = style.find("w:rPr", ns)
            if r_pr is None:
                r_pr = style.find(".//w:rPr", ns)
            if r_pr is None:
                r_pr = ET.SubElement(style, f"{W}rPr")
            color = r_pr.find("w:color", ns)
            if color is None:
                color = ET.SubElement(r_pr, f"{W}color")
            for attr in list(color.attrib):
                if attr != f"{W}val":
                    del color.attrib[attr]
            color.set(f"{W}val", "000000")
        return ET.tostring(root, encoding="UTF-8", xml_declaration=True)

    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        with zipfile.ZipFile(docx_path, "r") as zin:
            names = zin.namelist()
            zin.extractall(tmp_path)
        for fname in ("word/styles.xml", "word/stylesWithEffects.xml"):
            fpath = tmp_path / fname
            if fpath.is_file():
                fpath.write_bytes(patch(fpath.read_bytes()))
        backup = docx_path.with_suffix(".docx.bak")
        shutil.copy2(docx_path, backup)
        with zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for name in names:
                zout.write(tmp_path / name, name)


def body(doc: Document) -> None:
    doc.add_heading("Résumé", level=1)
    text(doc, RESUME)
    text(doc, "Mots-clés : " + MOTS_CLES)
    text(
        doc,
        "Note : générer la table des matières dans Word via Références → Table des matières → Mettre à jour.",
    )

    doc.add_heading("Partie I — Cahier des charges", level=1)

    doc.add_heading("1. Introduction et contexte", level=2)
    text(
        doc,
        "Le commerce local nécessite des outils numériques accessibles. Assigame est une marketplace "
        "permettant aux vendeurs de publier des annonces et aux visiteurs de consulter un catalogue "
        "public. Un administrateur valide les inscriptions et les publications avant mise en ligne.",
    )
    text(
        doc,
        "L'application est monolithique : Spring Boot sert l'API REST et les pages HTML statiques. "
        "L'authentification repose sur JWT ; les données sont stockées dans PostgreSQL.",
    )
    bullets(
        doc,
        [
            "Permettre aux vendeurs de publier et gérer des annonces produits.",
            "Offrir un catalogue public filtrable sans compte visiteur.",
            "Garantir la qualité via modération administrateur.",
        ],
    )

    doc.add_heading("2. Analyse du besoin", level=2)
    text(
        doc,
        "Problématique : comment permettre à des vendeurs locaux de publier des annonces tout en "
        "contrôlant la qualité des contenus ? La solution retenue impose une validation "
        "administrateur avant toute mise en ligne d'un compte vendeur ou d'un produit.",
    )
    text(
        doc,
        "Périmètre inclus : catalogue public, inscription et authentification vendeur, espace vendeur "
        "(dashboard, publication, gestion), back-office admin (modération, catégories, offres). "
        "Hors périmètre : panier d'achat, paiement en ligne réel, compte client acheteur et "
        "messagerie interne. L'écran de paiement d'abonnement vendeur est une simulation UI.",
    )

    doc.add_heading("3. Acteurs du système", level=2)
    table(doc, ["Acteur", "Description"], ACTEURS)

    doc.add_heading("4. Exigences fonctionnelles et non fonctionnelles", level=2)
    table(doc, ["ID", "Exigence fonctionnelle"], EXIGENCES_FONCTIONNELLES_COMPACT)
    table(doc, ["ID", "Exigence non fonctionnelle"], EXIGENCES_NON_FONCTIONNELLES)
    bullets(
        doc,
        [
            "Contraintes : Java 17, Spring Boot 4, PostgreSQL, frontend sans framework.",
            "Risques : secret JWT en dev, stockage local des images, paiement abonnement simulé.",
        ],
    )

    doc.add_heading("5. Cas d'utilisation", level=2)
    table(doc, ["ID", "Acteur", "Cas d'utilisation", "Résultat"], CAS_UTILISATION_COMPACT)
    figure(doc, DIAGRAMS_DIR / "use-case.png", "Figure F1 — Diagramme de cas d'utilisation", 5.5)

    doc.add_heading("6. Diagramme de classes", level=2)
    text(
        doc,
        "Quatre entités JPA : TypeUtilisateur, Utilisateur, CategorieProduit, Produit. "
        "Relations : TypeUtilisateur 1→* Utilisateur 1→* Produit ; CategorieProduit 1→* Produit. "
        "Statuts : ACTIF, EN_ATTENTE, REFUSE, SUSPENDU.",
    )
    figure(doc, DIAGRAMS_DIR / "class-diagram.png", "Figure F2 — Diagramme de classes", 4.8)

    doc.add_page_break()
    doc.add_heading("Partie II — Conception et réalisation", level=1)

    doc.add_heading("7. Architecture du système", level=2)
    text(
        doc,
        "Le navigateur accède aux pages statiques et à l'API REST. Spring Security valide les JWT. "
        "Les services métier utilisent JPA pour PostgreSQL ; les images sont dans uploads/produits/.",
    )
    figure(doc, DIAGRAMS_DIR / "architecture.png", "Figure F3 — Architecture", 5.2)
    table(doc, ["Couche", "Technologie"], STACK_TECHNIQUE)

    doc.add_heading("8. Règles métier, sécurité et API", level=2)
    bullets(
        doc,
        [
            "4 à 6 images obligatoires ; description max 2000 caractères.",
            "Quotas : Particulier 5, Professionnel 15, Partenaire Vip illimité (produits actifs).",
            "Modification vendeur → produit EN_ATTENTE ; seuls les ACTIF sont publics.",
            "JWT dans Authorization: Bearer ; tokens admin et vendeur séparés (localStorage).",
            "BCrypt pour les mots de passe ; vendeur ne modifie que ses propres produits.",
        ],
    )
    table(doc, ["Endpoint", "Accès", "Description"], API_ENDPOINTS_COMPACT)

    seq = DIAGRAMS_DIR / "sequence-moderation.png"
    if seq.is_file():
        figure(doc, seq, "Figure F4 — Séquence : modération d'un produit", 5.0)

    doc.add_heading("9. Interfaces utilisateur", level=2)
    text(doc, "Captures d'écran des principales interfaces de la plateforme Assigame.")

    # Trois écrans clés en pleine largeur (occupe plus d'espace)
    for fid, fname, title in SCREENSHOTS[:3]:
        figure(doc, SCREENSHOTS_DIR / fname, f"Figure {fid} — {title}", 5.5)

    doc.add_page_break()
    text(doc, "Espace vendeur et administration :")
    for i in range(3, len(SCREENSHOTS), 2):
        screenshots_pair(doc, SCREENSHOTS[i : i + 2], width=3.4)

    doc.add_page_break()
    doc.add_heading("Partie III — Validation et conclusion", level=1)

    doc.add_heading("10. Scénarios de test", level=2)
    table(doc, ["ID", "Scénario", "Résultat attendu"], SCENARIOS_TEST)

    doc.add_heading("11. Difficultés rencontrées", level=2)
    bullets(
        doc,
        [
            "Erreurs 403 admin : résolu en réordonnant les règles /api/admin/** dans SecurityConfig.",
            "Comptage catégories : calcul dynamique via l'API produits au lieu d'un tiret en dur.",
            "Uploads : validation 4–6 images et stockage CSV côté service métier.",
            "Sessions : séparation assigame_admin_token et assigame_vendor_token.",
        ],
    )

    doc.add_heading("12. Bilan, perspectives et conclusion", level=2)
    bullets(
        doc,
        [
            "Bilan : catalogue, parcours vendeur, modération et gestion catalogue livrés.",
            "Perspectives : paiement Stripe/Mobile Money, compte client, emails, déploiement cloud.",
        ],
    )
    text(
        doc,
        "Assigame démontre la conception d'une marketplace complète : backend Java, modélisation "
        "relationnelle, sécurité JWT et intégration front statique. La modération centralisée "
        "constitue le cœur fonctionnel du projet.",
    )

    doc.add_heading("Annexes", level=2)
    doc.add_heading("A. Arborescence du projet", level=3)
    tree = (
        "assigame/\n"
        "├── src/main/java/.../controller, service, entity, repository, security, config\n"
        "├── src/main/resources/static/ (HTML, CSS, JS — admin/, vendeur/)\n"
        "├── src/main/resources/application.properties\n"
        "└── uploads/produits/ (images uploadées)"
    )
    p = doc.add_paragraph()
    run = p.add_run(tree)
    run.font.name = "Consolas"
    run.font.size = Pt(9)

    doc.add_heading("B. Installation et comptes test", level=3)
    bullets(
        doc,
        [
            "Prérequis : Java 17, Maven 3.8+, PostgreSQL 14+.",
            "Créer la base : CREATE DATABASE basespring;",
            "Lancer : mvn spring-boot:run → http://localhost:8080",
            "Admin : admin@assigame.com / Admin1234",
        ],
    )

    doc.add_heading("C. Bibliographie", level=3)
    bullets(
        doc,
        [
            "Spring Boot — https://spring.io/projects/spring-boot",
            "Spring Security — https://docs.spring.io/spring-security/reference/",
            "PostgreSQL — https://www.postgresql.org/docs/",
            "JWT — https://jwt.io/",
        ],
    )


def page_count_via_libreoffice() -> int | None:
    import shutil
    import subprocess
    import tempfile

    if not shutil.which("libreoffice"):
        return None
    with tempfile.TemporaryDirectory() as tmp:
        pdf = Path(tmp) / "out.pdf"
        subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                tmp,
                str(OUTPUT),
            ],
            capture_output=True,
            check=False,
        )
        candidates = list(Path(tmp).glob("*.pdf"))
        if not candidates:
            return None
        result = subprocess.run(
            ["pdfinfo", str(candidates[0])],
            capture_output=True,
            text=True,
            check=False,
        )
        for line in result.stdout.splitlines():
            if line.startswith("Pages:"):
                return int(line.split(":")[1].strip())
    return None


def main() -> None:
    doc = Document()
    setup_styles(doc)
    cover(doc)
    body(doc)
    doc.save(str(OUTPUT))
    fix_heading_colors(OUTPUT)

    pages = page_count_via_libreoffice()
    print(f"Document généré : {OUTPUT}")
    print(f"Taille : {OUTPUT.stat().st_size / 1024:.1f} Ko")
    if pages:
        print(f"Pages : {pages} (objectif : {TARGET_PAGES})")


if __name__ == "__main__":
    main()
